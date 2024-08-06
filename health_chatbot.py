import os
from datetime import datetime
import streamlit as st
from streamlit_autorefresh import st_autorefresh
from langchain_groq import ChatGroq
from langchain_core.messages import SystemMessage, HumanMessage, ToolMessage
import threading
import time
from background import BackgroundCSSGenerator
from geopy.geocoders import Nominatim
from streamlit_folium import st_folium
import folium
import random
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
from googlesearch import search
from bs4 import BeautifulSoup
import requests


st.set_page_config(
    page_title="HealthMate",
    layout="wide",
    initial_sidebar_state="collapsed"
)

img1_path = r"C:\Users\praje\OneDrive\AppData\Desktop\project\health\work\Assests\black-background.gif"
img2_path = r"C:\Users\praje\OneDrive\AppData\Desktop\project\health\work\Assests\black-background.gif"
background_generator = BackgroundCSSGenerator(img1_path, img2_path)
page_bg_img = background_generator.generate_background_css()
st.markdown(page_bg_img, unsafe_allow_html=True)




# Set the environment variable for the API key
os.environ['GROQ_API_KEY'] = "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"

# Predefined responses for common symptoms
training_data = {
    "headache": "You might have a migraine or tension headache. Ensure to stay hydrated, rest in a dark and quiet room, and avoid bright lights. Over-the-counter pain relievers can help. Consult a healthcare professional if the pain persists or is severe.",
    "fever": "You might have a viral or bacterial infection. Rest, stay hydrated, and take fever-reducing medications like acetaminophen or ibuprofen. Monitor your symptoms and consult a doctor if the fever lasts more than a few days or is accompanied by other concerning symptoms.",
    "sore throat": "You might have tonsillitis, pharyngitis, or a mild viral infection. Gargle with warm salt water, drink plenty of fluids, and consider using throat lozenges. Consult a doctor if symptoms persist for more than a week or are accompanied by high fever.",
    "cough": "You might have a respiratory infection or allergies. Stay hydrated, use cough suppressants or expectorants if needed, and avoid irritants. Consult a healthcare professional if the cough lasts more than a few weeks or is accompanied by difficulty breathing.",
    "runny nose": "You might be experiencing allergies or a common cold. Use a saline nasal spray, stay hydrated, and consider antihistamines if allergies are suspected. Consult a doctor if symptoms are severe or persistent.",
    "nausea": "You might have a digestive issue, motion sickness, or an early sign of pregnancy. Drink clear fluids, rest, and avoid heavy or greasy foods. Consult a healthcare professional if nausea persists or is accompanied by other severe symptoms.",
    "dizziness": "You might have a vestibular issue, low blood pressure, or dehydration. Sit or lie down until the feeling passes, stay hydrated, and avoid sudden movements. Consult a doctor if dizziness is frequent, severe, or accompanied by other symptoms like fainting.",
    "chest pain": "You might be experiencing a range of issues from heartburn to more serious conditions like angina or a heart attack. Avoid exertion, monitor your symptoms, and seek immediate medical attention if the pain is severe, persistent, or accompanied by other symptoms like shortness of breath.",
    "back pain": "You might have muscle strain, poor posture, or a herniated disc. Rest, use heat or cold packs, and maintain good posture. Consult a healthcare professional if the pain persists, is severe, or is accompanied by other symptoms like numbness or weakness.",
    "abdominal pain": "You might have digestive issues, gastritis, or other gastrointestinal conditions. Rest, avoid spicy or heavy foods, and consider over-the-counter medications. Consult a doctor if the pain is severe, persistent, or accompanied by other symptoms like vomiting or changes in bowel habits.",
    "fatigue": "You might be experiencing stress, sleep deprivation, or an underlying health condition. Ensure you are getting adequate rest, managing stress, and eating a balanced diet. Consult a healthcare professional if fatigue persists despite lifestyle changes or is accompanied by other concerning symptoms.",
    "skin rash": "You might have an allergic reaction, eczema, or another dermatological issue. Avoid known allergens, use gentle skin care products, and consider over-the-counter anti-itch creams. Consult a doctor if the rash is severe, spreading, or accompanied by other symptoms like fever.",
    "shortness of breath": "You might have a respiratory infection, asthma, or other lung conditions. Ensure proper ventilation, use prescribed inhalers if applicable, and avoid strenuous activities. Seek immediate medical attention if shortness of breath is severe or accompanied by chest pain or dizziness.",
    "joint pain": "You might be experiencing arthritis, strain, or injury. Rest the affected joint, use ice or heat as needed, and consider over-the-counter pain relievers. Consult a healthcare professional if the pain persists, is severe, or is accompanied by swelling or stiffness.",
    "swollen legs": "You might have issues related to circulation, injury, or fluid retention. Elevate your legs, reduce salt intake, and avoid prolonged periods of standing. Consult a doctor if swelling persists, is severe, or is accompanied by other symptoms like pain or redness.",
    "sleep problems": "You might be experiencing insomnia or sleep disorders. Establish a regular sleep routine, create a comfortable sleep environment, and avoid stimulants before bedtime. Consult a healthcare professional if sleep issues persist or are affecting your daily life.",
    "earache": "You might have an ear infection or earwax buildup. Avoid inserting objects into your ear, use over-the-counter pain relievers, and consult a doctor if the pain is severe or persists for more than a few days.",
    "numbness": "You might be experiencing nerve compression, poor circulation, or a neurological issue. Rest the affected area, avoid repetitive motions, and consult a healthcare professional if numbness is persistent, severe, or accompanied by weakness.",
    "coughing up blood": "This could be a sign of a serious condition such as a respiratory infection, lung disease, or bleeding disorder. Seek immediate medical attention if you are coughing up blood or if it persists.",
    "weight loss": "You might have an underlying health condition, stress, or changes in diet. Ensure you are eating a balanced diet and consult a healthcare professional if weight loss is significant or unintentional.",
    "weight gain": "You might be experiencing changes in diet, lack of physical activity, or other underlying health issues. Consider evaluating your diet and exercise habits, and consult a healthcare professional if weight gain is significant or unexplained.",
    "frequent urination": "You might have a urinary tract infection, diabetes, or other conditions affecting the urinary system. Stay hydrated, avoid irritants, and consult a doctor if frequent urination is persistent or accompanied by other symptoms.",
    "painful urination": "You might have a urinary tract infection or other urinary issues. Increase fluid intake, avoid irritants, and consult a healthcare professional if painful urination persists or is severe.",
    "constipation": "You might be experiencing a lack of dietary fiber, dehydration, or other digestive issues. Increase fiber intake, drink plenty of water, and consider over-the-counter laxatives if needed. Consult a doctor if constipation is persistent or severe.",
    "diarrhea": "You might have a viral or bacterial infection, food poisoning, or digestive issues. Stay hydrated, avoid irritants, and consult a healthcare professional if diarrhea is severe, persistent, or accompanied by other symptoms.",
    "abdominal bloating": "You might have gas, indigestion, or other gastrointestinal issues. Avoid carbonated beverages, eat smaller meals, and consult a doctor if bloating is severe or persistent.",
    "drowsiness": "You might be experiencing poor sleep quality, fatigue, or medication side effects. Ensure you are getting adequate sleep and consult a healthcare professional if drowsiness is persistent or interfering with daily activities.",
    "itchy eyes": "You might have allergies or dry eyes. Use antihistamine eye drops, avoid rubbing your eyes, and consult a healthcare professional if the itching is severe or persistent.",
    "dry mouth": "You might be experiencing dehydration, medication side effects, or a salivary gland issue. Increase fluid intake, use sugar-free lozenges, and consult a doctor if dry mouth is persistent or severe.",
    "bloody stools": "This could be a sign of a serious condition such as gastrointestinal bleeding or hemorrhoids. Seek immediate medical attention if you notice blood in your stools or if symptoms persist.",
    "chronic cough": "You might have chronic bronchitis, asthma, or other respiratory conditions. Avoid irritants, use prescribed medications if applicable, and consult a healthcare professional if the cough persists for more than eight weeks.",
    "tremors": "You might have a neurological condition, anxiety, or medication side effects. Consult a healthcare professional for a proper evaluation if tremors are persistent or affecting daily life.",
    "hot flashes": "You might be experiencing hormonal changes related to menopause or other conditions. Wear breathable clothing, stay hydrated, and consult a healthcare professional if hot flashes are severe or frequent.",
    "cold sores": "You might have a herpes simplex virus infection. Use antiviral medications, avoid touching the sores, and consult a healthcare professional if outbreaks are frequent or severe.",
    "sweating excessively": "You might have a condition such as hyperhidrosis or be experiencing stress or anxiety. Use antiperspirants, stay hydrated, and consult a healthcare professional if excessive sweating is persistent or interfering with daily activities.",
    "muscle cramps": "You might be experiencing dehydration, overuse of muscles, or electrolyte imbalances. Stay hydrated, stretch the affected muscle, and consult a doctor if cramps are frequent or severe.",
    "hair loss": "You might be experiencing stress, hormonal changes, or nutritional deficiencies. Ensure you are getting adequate nutrients and consult a healthcare professional if hair loss is significant or persistent.",
    "blurry vision": "You might have eye strain, refractive errors, or other vision issues. Rest your eyes, use appropriate corrective lenses if needed, and consult an eye specialist if blurry vision persists or worsens.",
    "itchy skin": "You might have dry skin, allergies, or eczema. Use gentle skin care products, avoid irritants, and consult a healthcare professional if itching is severe or persistent.",
    "feeling faint": "You might be experiencing low blood pressure, dehydration, or other underlying issues. Sit or lie down, stay hydrated, and consult a doctor if fainting episodes are frequent or severe.",
    "high blood pressure": "You might be experiencing hypertension, which can be managed with lifestyle changes and medications. Monitor your blood pressure regularly, reduce salt intake, and consult a healthcare professional for personalized advice.",
    "low blood pressure": "You might be experiencing hypotension due to dehydration, medication side effects, or other conditions. Increase fluid intake, avoid standing up quickly, and consult a healthcare professional if symptoms persist or are severe.",
    "tingling sensation": "You might have nerve compression, poor circulation, or other neurological issues. Avoid repetitive motions, rest the affected area, and consult a healthcare professional if tingling is persistent or severe.",
    "painful muscles": "You might be experiencing muscle strain, overuse, or an underlying condition. Rest, apply heat or cold packs, and consult a doctor if muscle pain is severe or persistent.",
    "joint stiffness": "You might have arthritis, overuse, or an injury. Maintain an active lifestyle, use over-the-counter pain relievers, and consult a healthcare professional if stiffness is persistent or affecting daily activities.",
    "ear ringing": "You might have tinnitus, which could be caused by exposure to loud noises or other factors. Avoid loud environments, use white noise machines, and consult a healthcare professional if ringing is persistent or bothersome.",
    "bad breath": "You might have poor oral hygiene, gum disease, or digestive issues. Maintain good oral hygiene, drink plenty of water, and consult a dentist if bad breath persists despite regular brushing and flossing.",
    "sweaty palms": "You might have hyperhidrosis or be experiencing stress. Use antiperspirants designed for palms, practice stress-relief techniques, and consult a healthcare professional if sweating is excessive or interfering with daily life.",
    "nosebleeds": "You might have dry nasal passages, allergies, or other underlying conditions. Use a saline nasal spray, avoid picking your nose, and consult a healthcare professional if nosebleeds are frequent or severe.",
    "persistent hiccups": "You might have a minor irritation of the diaphragm or other underlying issues. Try holding your breath, drinking a glass of water quickly, and consult a doctor if hiccups persist for more than 48 hours.",
    "muscle weakness": "You might be experiencing fatigue, neurological issues, or muscle disorders. Ensure adequate rest and nutrition, and consult a healthcare professional if weakness is persistent or affecting daily activities.",
    "frequent headaches": "You might be experiencing migraines, tension headaches, or other underlying conditions. Keep a headache diary, manage stress, and consult a healthcare professional if headaches are frequent or severe.",
    "chronic fatigue": "You might have an underlying health condition, sleep disorder, or chronic fatigue syndrome. Evaluate your sleep quality, manage stress, and consult a healthcare professional for a thorough evaluation if fatigue is persistent or debilitating.",
    "blurred vision": "You might have eye strain, refractive errors, or other eye conditions. Rest your eyes, use corrective lenses if needed, and consult an eye specialist if blurred vision persists or worsens.",
    "painful swallowing": "You might have an infection, inflammation, or other throat issues. Gargle with warm salt water, avoid irritants, and consult a healthcare professional if painful swallowing is persistent or severe.",
    "sudden weight loss": "You might have an underlying health condition such as diabetes, thyroid disorders, or gastrointestinal issues. Ensure you are eating a balanced diet and consult a healthcare professional if weight loss is significant or unexplained.",
    "frequent urination at night": "You might have a urinary tract infection, diabetes, or other conditions. Reduce fluid intake before bedtime and consult a healthcare professional if nocturia is persistent or affecting your quality of life.",
    "dry skin": "You might be experiencing dehydration, environmental factors, or skin conditions. Use moisturizers, stay hydrated, and consult a dermatologist if dry skin persists or is severe.",
    "itchy scalp": "You might have dandruff, scalp psoriasis, or other scalp conditions. Use gentle shampoos, avoid scratching, and consult a dermatologist if itching is persistent or accompanied by other symptoms.",
    "swollen joints": "You might have arthritis, injury, or fluid retention. Rest the affected joint, apply ice, and consult a healthcare professional if swelling is persistent or accompanied by pain or redness.",
    "tender abdomen": "You might have digestive issues, gastritis, or other gastrointestinal conditions. Rest, avoid heavy or spicy foods, and consult a doctor if tenderness is persistent or accompanied by other symptoms.",
    "persistent cough": "You might have chronic bronchitis, asthma, or other respiratory issues. Avoid irritants, use prescribed medications, and consult a healthcare professional if the cough persists for more than eight weeks.",
    "loss of appetite": "You might have an underlying health condition, stress, or medication side effects. Ensure you are eating a balanced diet and consult a healthcare professional if loss of appetite is significant or persistent.",
    "night sweats": "You might be experiencing hormonal changes, infections, or other underlying issues. Use breathable bedding, manage stress, and consult a healthcare professional if night sweats are frequent or severe.",
    "painful menstruation": "You might have dysmenorrhea or other menstrual issues. Use pain relievers, apply heat to the lower abdomen, and consult a healthcare professional if menstrual pain is severe or affecting daily life.",
    "persistent sneezing": "You might have allergies or a respiratory infection. Avoid allergens, use antihistamines, and consult a healthcare professional if sneezing is frequent or severe.",
    "abdominal cramping": "You might have digestive issues, menstrual cramps, or other conditions. Rest, use heat packs, and consult a healthcare professional if cramping is severe or persistent.",
    "feeling anxious": "You might be experiencing stress, anxiety disorder, or other mental health issues. Practice relaxation techniques, seek support from friends or a mental health professional, and consult a healthcare professional if anxiety is persistent or interfering with daily life.",
    "irregular heartbeat": "You might have arrhythmia or other heart conditions. Monitor your symptoms, avoid stimulants, and consult a healthcare professional for a thorough evaluation if irregular heartbeat is persistent or severe.",
    "hoarseness": "You might have laryngitis, vocal strain, or other throat issues. Rest your voice, avoid irritants, and consult a healthcare professional if hoarseness persists for more than a few weeks.",
    "excessive thirst": "You might have diabetes, dehydration, or other conditions. Drink plenty of fluids and consult a healthcare professional if excessive thirst is persistent or accompanied by other symptoms.",
    "difficulty concentrating": "You might be experiencing fatigue, stress, or cognitive issues. Ensure adequate rest, manage stress, and consult a healthcare professional if concentration difficulties are persistent or affecting daily life.",
    "bruising easily": "You might have a bleeding disorder, vitamin deficiency, or other issues. Avoid activities that cause injury, ensure you are getting adequate nutrition, and consult a healthcare professional if bruising is frequent or severe.",
    "bladder pain": "You might have a urinary tract infection, bladder inflammation, or other issues. Stay hydrated, avoid irritants, and consult a healthcare professional if bladder pain is persistent or severe.",
    "painful muscles": "You might be experiencing muscle strain, overuse, or other conditions. Rest, use ice or heat as needed, and consult a healthcare professional if muscle pain is persistent or affecting daily activities.",
    "red or swollen eyes": "You might have conjunctivitis, allergies, or other eye issues. Use cool compresses, avoid irritants, and consult an eye specialist if redness or swelling persists or is accompanied by vision changes.",
    "skin ulcers": "You might have pressure sores, infections, or other skin conditions. Keep the affected area clean and dry, avoid pressure, and consult a healthcare professional if ulcers are severe or persistent.",
    "gastrointestinal bleeding": "You might have a serious condition such as ulcers or hemorrhoids. Seek immediate medical attention if you notice blood in your stool or vomit, or if you experience severe abdominal pain.",
    "sweating at night": "You might have hormonal changes, infections, or other underlying conditions. Use breathable bedding, stay hydrated, and consult a healthcare professional if night sweats are frequent or severe.",
    "bleeding gums": "You might have gum disease, vitamin deficiencies, or other issues. Practice good oral hygiene, avoid irritating foods, and consult a dentist if bleeding persists or is severe.",
    "painful joints": "You might have arthritis, injury, or inflammation. Rest the affected joint, use pain relievers, and consult a healthcare professional if the pain is severe or persistent.",
    "difficulty breathing": "You might have a respiratory infection, asthma, or other lung issues. Ensure proper ventilation, use prescribed inhalers if applicable, and seek immediate medical attention if difficulty breathing is severe or persistent.",
    "excessive urination": "You might have diabetes, a urinary tract infection, or other conditions. Increase fluid intake, monitor your symptoms, and consult a healthcare professional if excessive urination is persistent or affecting your quality of life.",
    "itchy throat": "You might have allergies, a throat infection, or irritation. Use antihistamines, stay hydrated, and consult a healthcare professional if itching is persistent or severe.",
    "painful breathing": "You might have a respiratory infection, pleuritis, or other lung conditions. Rest, avoid irritants, and seek medical attention if pain worsens or persists.",
    "persistent dry cough": "You might have a chronic respiratory condition or allergies. Use prescribed medications if applicable, stay hydrated, and consult a healthcare professional if the cough persists or worsens.",
    "painful or burning urination": "You might have a urinary tract infection or bladder inflammation. Increase fluid intake, use over-the-counter pain relief, and consult a healthcare professional if symptoms are severe or persistent.",
    "muscle soreness": "You might have overexerted yourself or be experiencing a muscle strain. Rest, apply heat or cold packs as needed, and consult a healthcare professional if soreness is severe or persistent.",
    "fatigue after exercise": "You might have overexerted yourself or not properly hydrated. Rest, hydrate, and ensure you're following a balanced exercise routine. Consult a healthcare professional if fatigue is severe or lasts longer than expected.",
    "painful or swollen hands": "You might have arthritis, repetitive strain injuries, or other conditions. Rest your hands, use ice or heat as needed, and consult a healthcare professional if symptoms persist or worsen.",
    "difficulty sleeping": "You might be experiencing stress, insomnia, or other sleep disorders. Establish a regular sleep routine, create a restful environment, and consult a healthcare professional if sleep difficulties persist or impact daily life.",
    "dry or flaky skin": "You might have a skin condition or be experiencing environmental dryness. Use moisturizing lotions, stay hydrated, and consult a dermatologist if dryness is severe or persistent.",
    "redness around the eyes": "You might have conjunctivitis, allergies, or other eye issues. Use cool compresses, avoid irritants, and consult an eye specialist if redness persists or is accompanied by other symptoms.",
    "painful or swollen feet": "You might have an injury, fluid retention, or circulatory issues. Rest your feet, elevate them, and consult a healthcare professional if pain or swelling is severe or persistent.",
    "persistent bloating": "You might have digestive issues, food intolerances, or other gastrointestinal conditions. Avoid carbonated beverages, eat smaller meals, and consult a healthcare professional if bloating is severe or persistent.",
    "painful or swollen lymph nodes": "You might have an infection, inflammation, or other underlying condition. Rest, stay hydrated, and consult a healthcare professional if lymph node pain or swelling persists or worsens.",
    "painful urination at night": "You might have a urinary tract infection or other urinary issues. Increase fluid intake during the day and consult a healthcare professional if symptoms persist or are severe.",
    "frequent mood swings": "You might be experiencing hormonal changes, stress, or mental health issues. Manage stress, seek support if needed, and consult a healthcare professional if mood swings are severe or impacting daily life.",
    "pain in the chest": "You might be experiencing heartburn, angina, or other serious conditions. Avoid heavy meals, monitor your symptoms, and seek immediate medical attention if chest pain is severe or persistent.",
    "persistent ringing in the ears": "You might have tinnitus, often related to exposure to loud noises or other factors. Avoid loud environments, use white noise machines, and consult a healthcare professional if ringing is persistent or bothersome.",
    "frequent nosebleeds": "You might have dry nasal passages, allergies, or other conditions. Use a saline nasal spray, avoid picking your nose, and consult a healthcare professional if nosebleeds are frequent or severe.",
    "persistent sore throat": "You might have a throat infection, allergies, or irritation. Gargle with warm salt water, use throat lozenges, and consult a healthcare professional if the sore throat persists for more than a week.",
    "difficulty swallowing": "You might have an infection, inflammation, or other throat issues. Avoid irritants, eat soft foods, and consult a healthcare professional if swallowing difficulties are severe or persistent.",
    "painful or swollen abdomen": "You might have digestive issues, inflammation, or other gastrointestinal conditions. Rest, avoid heavy foods, and consult a healthcare professional if pain or swelling persists or is severe.",
    "painful joints or muscles": "You might have arthritis, muscle strain, or other conditions. Rest the affected area, use heat or cold packs, and consult a healthcare professional if pain is severe or persistent.",
    "persistent cough or sore throat": "You might have a respiratory infection, allergies, or other issues. Rest, stay hydrated, and consult a healthcare professional if symptoms are severe or persist for more than a week.",
    "persistent headache": "You might have a migraine, tension headache, or other issues. Stay hydrated, rest, and avoid triggers. Consult a healthcare professional if headaches are severe or frequent.",
    "painful or swollen throat": "You might have a throat infection, tonsillitis, or other issues. Gargle with warm salt water, stay hydrated, and consult a healthcare professional if pain is severe or persistent.",
    "feeling shaky": "You might be experiencing low blood sugar, anxiety, or other issues. Eat a balanced meal, manage stress, and consult a healthcare professional if shaking is persistent or severe.",
    "difficult breathing": "You might have asthma, a respiratory infection, or other lung issues. Use prescribed inhalers if applicable, avoid irritants, and seek medical attention if breathing difficulties are severe or persistent.",
    "painful or swollen joints": "You might have arthritis, injury, or inflammation. Rest the affected joint, use ice or heat, and consult a healthcare professional if pain or swelling persists or is severe.",
    "painful swallowing or sore throat": "You might have an infection or irritation. Gargle with warm salt water, drink fluids, and consult a healthcare professional if symptoms are severe or persistent.",
    "constant fatigue": "You might have an underlying health condition, sleep issues, or chronic fatigue syndrome. Ensure you are getting enough rest and consult a healthcare professional if fatigue is persistent or severe.",
    "dry cough or sore throat": "You might have a respiratory infection, allergies, or irritation. Stay hydrated, use cough remedies, and consult a healthcare professional if symptoms are severe or persist for more than a week.",
    "abdominal discomfort": "You might have digestive issues, gas, or other gastrointestinal problems. Rest, eat smaller meals, and consult a healthcare professional if discomfort persists or worsens.",
    "painful urination or frequent urination": "You might have a urinary tract infection or other issues. Increase fluid intake, use pain relief as needed, and consult a healthcare professional if symptoms are severe or persistent.",
    "persistent stomach pain": "You might have digestive issues, ulcers, or other conditions. Rest, avoid irritants, and consult a healthcare professional if pain is severe or persistent.",
    "painful menstruation": "You might have dysmenorrhea or other menstrual issues. Use pain relievers, apply heat to the abdomen, and consult a healthcare professional if menstrual pain is severe or affecting daily activities.",
    "chronic dry skin": "You might have a skin condition or be experiencing environmental dryness. Use moisturizers, stay hydrated, and consult a dermatologist if dryness is severe or persistent.",
    "itchy or red eyes": "You might have allergies, conjunctivitis, or other eye issues. Use cool compresses, avoid irritants, and consult an eye specialist if symptoms are severe or persistent.",
    "persistent back pain": "You might have muscle strain, poor posture, or a spinal condition. Rest, maintain good posture, and consult a healthcare professional if pain is severe or persistent.",
    "swollen feet or ankles": "You might have fluid retention, injury, or circulatory issues. Elevate your feet, reduce salt intake, and consult a healthcare professional if swelling is severe or persistent.",
    "dry cough or sore throat": "You might have a respiratory infection, allergies, or irritation. Stay hydrated, use cough remedies, and consult a healthcare professional if symptoms are severe or persist for more than a week.",
    "frequent headaches": "You might have tension headaches, migraines, or other issues. Keep a headache diary, manage stress, and consult a healthcare professional if headaches are frequent or severe.",
    "painful or swollen throat": "You might have a throat infection, tonsillitis, or other issues. Gargle with warm salt water, stay hydrated, and consult a healthcare professional if symptoms are severe or persistent.",
    "feeling lightheaded": "You might have low blood pressure, dehydration, or other issues. Sit or lie down, stay hydrated, and consult a healthcare professional if lightheadedness is persistent or severe.",
    "dry mouth or throat": "You might have dehydration, medication side effects, or a dry environment. Increase fluid intake, use saliva substitutes, and consult a healthcare professional if symptoms are persistent or severe.",
    "painful or swollen abdomen": "You might have digestive issues, inflammation, or other conditions. Rest, avoid heavy foods, and consult a healthcare professional if symptoms are severe or persistent.",
    "chronic fatigue": "You might have an underlying health condition, sleep disorder, or chronic fatigue syndrome. Ensure you are getting adequate rest and consult a healthcare professional if fatigue is persistent or debilitating.",
    "painful urination": "You might have a urinary tract infection or other urinary issues. Increase fluid intake, use over-the-counter pain relief, and consult a healthcare professional if symptoms are severe or persistent.",
    "persistent joint pain": "You might have arthritis, injury, or inflammation. Rest the affected joint, use pain relievers, and consult a healthcare professional if pain is severe or persistent.",
    "sweating excessively": "You might have hyperhidrosis or be experiencing stress. Use antiperspirants designed for excessive sweating, practice stress-relief techniques, and consult a healthcare professional if sweating is severe or interfering with daily life.",
    "painful muscles": "You might be experiencing muscle strain, overuse, or other conditions. Rest the affected muscles, apply heat or cold as needed, and consult a healthcare professional if pain is severe or persistent.",
    "chronic headache": "You might have tension headaches, migraines, or other issues. Stay hydrated, manage stress, and consult a healthcare professional if headaches are frequent or severe.",
    "painful or swollen joints": "You might have arthritis, injury, or inflammation. Rest the affected joints, use ice or heat as needed, and consult a healthcare professional if symptoms are severe or persistent.",
    "persistent cough": "You might have a respiratory infection, allergies, or other conditions. Use cough remedies, stay hydrated, and consult a healthcare professional if symptoms persist for more than a week.",
    "persistent or severe headache": "You might be experiencing a migraine, tension headache, or other condition. Keep a headache diary, manage stress, and consult a healthcare professional if headaches are severe or frequent.",
    "painful swallowing": "You might have an infection, inflammation, or other throat issues. Gargle with warm salt water, avoid irritants, and consult a healthcare professional if swallowing difficulties are severe or persistent.",
    "dry throat or mouth": "You might have dehydration, medication side effects, or other issues. Increase fluid intake, use saliva substitutes, and consult a healthcare professional if symptoms are persistent or severe.",
    "painful breathing": "You might have a respiratory infection, pleuritis, or other lung conditions. Rest, avoid irritants, and consult a healthcare professional if pain worsens or persists.",
    "painful or burning urination": "You might have a urinary tract infection or other conditions. Increase fluid intake, use over-the-counter pain relief, and consult a healthcare professional if symptoms are severe or persistent.",
    "swollen lymph nodes": "You might have an infection, inflammation, or other underlying condition. Rest, stay hydrated, and consult a healthcare professional if swelling persists or worsens.",
    "painful abdomen": "You might have digestive issues, inflammation, or other conditions. Rest, avoid heavy foods, and consult a healthcare professional if pain is severe or persistent.",
    "painful or swollen throat": "You might have a throat infection, tonsillitis, or other issues. Gargle with warm salt water, stay hydrated, and consult a healthcare professional if pain is severe or persistent.",
    "painful menstruation": "You might have dysmenorrhea or other menstrual issues. Use pain relievers, apply heat to the lower abdomen, and consult a healthcare professional if pain is severe or affecting daily life.",
    "sore throat or dry cough": "You might have a respiratory infection, allergies, or irritation. Stay hydrated, use throat lozenges, and consult a healthcare professional if symptoms are severe or persistent.",
    "swollen feet or ankles": "You might have fluid retention, injury, or circulatory issues. Elevate your feet, reduce salt intake, and consult a healthcare professional if swelling is severe or persistent.",
    "chronic dry skin": "You might have a skin condition or be experiencing environmental dryness. Use moisturizers, stay hydrated, and consult a dermatologist if dryness is severe or persistent.",
    "itchy scalp": "You might have dandruff, psoriasis, or other scalp conditions. Use gentle shampoos, avoid scratching, and consult a dermatologist if itching is persistent or accompanied by other symptoms.",
    "painful urination": "You might have a urinary tract infection or other conditions. Increase fluid intake, use over-the-counter pain relief, and consult a healthcare professional if symptoms are severe or persistent.",
    "painful or swollen feet": "You might have an injury, fluid retention, or circulatory issues. Rest, elevate your feet, and consult a healthcare professional if symptoms persist or worsen.",
    "painful or swollen abdomen": "You might have digestive issues, inflammation, or other conditions. Rest, avoid heavy foods, and consult a healthcare professional if symptoms are severe or persistent.",
    "persistent headache or migraine": "You might have tension headaches, migraines, or other issues. Stay hydrated, manage stress, and consult a healthcare professional if headaches are frequent or severe.",
    "painful breathing": "You might have a respiratory infection, pleuritis, or other lung conditions. Rest, avoid irritants, and consult a healthcare professional if pain worsens or persists.",
    "painful urination": "You might have a urinary tract infection or other conditions. Increase fluid intake, use over-the-counter pain relief, and consult a healthcare professional if symptoms are severe or persistent.",
    "painful joints": "You might have arthritis, injury, or inflammation. Rest the affected joints, use ice or heat as needed, and consult a healthcare professional if pain is severe or persistent.",
    "chronic dry skin": "You might have a skin condition or be experiencing environmental dryness. Use moisturizers, stay hydrated, and consult a dermatologist if dryness is severe or persistent.",
    "persistent cough": "You might have a respiratory infection, allergies, or other conditions. Use cough remedies, stay hydrated, and consult a healthcare professional if symptoms persist for more than a week.",
    "painful or swollen abdomen": "You might have digestive issues, inflammation, or other conditions. Rest, avoid heavy foods, and consult a healthcare professional if pain is severe or persistent.",
    "persistent or severe headache": "You might be experiencing a migraine, tension headache, or other condition. Keep a headache diary, manage stress, and consult a healthcare professional if headaches are severe or frequent.",
    "painful or swollen feet": "You might have an injury, fluid retention, or circulatory issues. Rest your feet, elevate them, and consult a healthcare professional if symptoms persist or worsen.",
    "painful or swollen throat": "You might have a throat infection, tonsillitis, or other issues. Gargle with warm salt water, stay hydrated, and consult a healthcare professional if symptoms are severe or persistent.",
    "persistent headache or migraine": "You might have a tension headache, migraine, or other issue. Stay hydrated, manage stress, and consult a healthcare professional if headaches are frequent or severe.",
    "painful or swollen joints": "You might have arthritis, injury, or inflammation. Rest the affected joints, use ice or heat as needed, and consult a healthcare professional if pain is severe or persistent.",
    "dry mouth or throat": "You might have dehydration, medication side effects, or other issues. Increase fluid intake, use saliva substitutes, and consult a healthcare professional if symptoms are persistent or severe.",
    "painful urination": "You might have a urinary tract infection or other urinary issues. Increase fluid intake, use over-the-counter pain relief, and consult a healthcare professional if symptoms are severe or persistent.",
    "persistent cough or sore throat": "You might have a respiratory infection, allergies, or other issues. Stay hydrated, use cough remedies, and consult a healthcare professional if symptoms are severe or persist for more than a week.",
    "painful or swollen abdomen": "You might have digestive issues, inflammation, or other conditions. Rest, avoid heavy foods, and consult a healthcare professional if pain is severe or persistent.",
    "persistent headaches or migraines": "You might have tension headaches, migraines, or other issues. Keep a headache diary, manage stress, and consult a healthcare professional if headaches are frequent or severe.",
    "painful or swollen throat": "You might have a throat infection, tonsillitis, or other issues. Gargle with warm salt water, stay hydrated, and consult a healthcare professional if symptoms are severe or persistent.",
    "persistent or severe headache": "You might be experiencing a migraine, tension headache, or other condition. Keep a headache diary, manage stress, and consult a healthcare professional if headaches are severe or frequent.",
    "painful or swollen feet": "You might have an injury, fluid retention, or circulatory issues. Rest your feet, elevate them, and consult a healthcare professional if symptoms persist or worsen.",
    "painful or swollen throat": "You might have a throat infection, tonsillitis, or other issues. Gargle with warm salt water, stay hydrated, and consult a healthcare professional if symptoms are severe or persistent.",
    "chronic dry skin": "You might have a skin condition or be experiencing environmental dryness. Use moisturizers, stay hydrated, and consult a dermatologist if dryness is severe or persistent.",
    "dry cough or sore throat": "You might have a respiratory infection, allergies, or irritation. Stay hydrated, use cough remedies, and consult a healthcare professional if symptoms are severe or persist for more than a week.",
    "painful urination or frequent urination": "You might have a urinary tract infection or other issues. Increase fluid intake, use pain relief as needed, and consult a healthcare professional if symptoms are severe or persistent.",
    "painful menstruation": "You might have dysmenorrhea or other menstrual issues. Use pain relievers, apply heat to the abdomen, and consult a healthcare professional if menstrual pain is severe or affecting daily life.",
    "painful or swollen joints": "You might have arthritis, injury, or inflammation. Rest the affected joints, use ice or heat as needed, and consult a healthcare professional if pain is severe or persistent.",
    "persistent headaches or migraines": "You might have tension headaches, migraines, or other issues. Keep a headache diary, manage stress, and consult a healthcare professional if headaches are frequent or severe.",
    "persistent cough or sore throat": "You might have a respiratory infection, allergies, or other issues. Stay hydrated, use cough remedies, and consult a healthcare professional if symptoms are severe or persist for more than a week.",
    "painful or swollen feet": "You might have an injury, fluid retention, or circulatory issues. Rest your feet, elevate them, and consult a healthcare professional if symptoms persist or worsen.",
    "painful urination": "You might have a urinary tract infection or other urinary issues. Increase fluid intake, use over-the-counter pain relief, and consult a healthcare professional if symptoms are severe or persistent.",
    "persistent dry skin": "You might have a skin condition or be experiencing environmental dryness. Use moisturizers, stay hydrated, and consult a dermatologist if dryness is severe or persistent.",
    "painful or swollen throat": "You might have a throat infection, tonsillitis, or other issues. Gargle with warm salt water, stay hydrated, and consult a healthcare professional if symptoms are severe or persistent.",
    "persistent headache or migraine": "You might have tension headaches, migraines, or other issues. Keep a headache diary, manage stress, and consult a healthcare professional if headaches are frequent or severe.",
    "painful or swollen abdomen": "You might have digestive issues, inflammation, or other conditions. Rest, avoid heavy foods, and consult a healthcare professional if symptoms are severe or persistent.",
    "painful or swollen feet": "You might have an injury, fluid retention, or circulatory issues. Rest your feet, elevate them, and consult a healthcare professional if symptoms persist or worsen.",
    "painful or swollen throat": "You might have a throat infection, tonsillitis, or other issues. Gargle with warm salt water, stay hydrated, and consult a healthcare professional if symptoms are severe or persistent.",
    "persistent or severe headache": "You might be experiencing a migraine, tension headache, or other condition. Keep a headache diary, manage stress, and consult a healthcare professional if headaches are severe or frequent.",
    "painful or swollen joints": "You might have arthritis, injury, or inflammation. Rest the affected joints, use ice or heat as needed, and consult a healthcare professional if pain is severe or persistent.",
    "painful urination": "You might have a urinary tract infection or other urinary issues. Increase fluid intake, use over-the-counter pain relief, and consult a healthcare professional if symptoms are severe or persistent.",
    "chronic dry skin": "You might have a skin condition or be experiencing environmental dryness. Use moisturizers, stay hydrated, and consult a dermatologist if dryness is severe or persistent.",
    "painful or swollen throat": "You might have a throat infection, tonsillitis, or other issues. Gargle with warm salt water, stay hydrated, and consult a healthcare professional if symptoms are severe or persistent.",
    "persistent headaches or migraines": "You might have tension headaches, migraines, or other issues. Keep a headache diary, manage stress, and consult a healthcare professional if headaches are frequent or severe.",
    "painful or swollen feet": "You might have an injury, fluid retention, or circulatory issues. Rest your feet, elevate them, and consult a healthcare professional if symptoms persist or worsen.",
    "painful urination": "You might have a urinary tract infection or other urinary issues. Increase fluid intake, use over-the-counter pain relief, and consult a healthcare professional if symptoms are severe or persistent.",
    "persistent dry skin": "You might have a skin condition or be experiencing environmental dryness. Use moisturizers, stay hydrated, and consult a dermatologist if dryness is severe or persistent.",
    "painful or swollen throat": "You might have a throat infection, tonsillitis, or other issues. Gargle with warm salt water, stay hydrated, and consult a healthcare professional if symptoms are severe or persistent.",
    "painful or swollen abdomen": "You might have digestive issues, inflammation, or other conditions. Rest, avoid heavy foods, and consult a healthcare professional if symptoms are severe or persistent.",
    "persistent headache or migraine": "You might be experiencing a migraine, tension headache, or other condition. Keep a headache diary, manage stress, and consult a healthcare professional if headaches are severe or frequent.",
    "painful or swollen joints": "You might have arthritis, injury, or inflammation. Rest the affected joints, use ice or heat as needed, and consult a healthcare professional if pain is severe or persistent.",
    "painful urination": "You might have a urinary tract infection or other urinary issues. Increase fluid intake, use over-the-counter pain relief, and consult a healthcare professional if symptoms are severe or persistent.",
    "painful or swollen throat": "You might have a throat infection, tonsillitis, or other issues. Gargle with warm salt water, stay hydrated, and consult a healthcare professional if symptoms are severe or persistent.",
    "chronic dry skin": "You might have a skin condition or be experiencing environmental dryness. Use moisturizers, stay hydrated, and consult a dermatologist if dryness is severe or persistent.",
    "painful urination or frequent urination": "You might have a urinary tract infection or other issues. Increase fluid intake, use pain relief as needed, and consult a healthcare professional if symptoms are severe or persistent.",
    "painful menstruation": "You might have dysmenorrhea or other menstrual issues. Use pain relievers, apply heat to the abdomen, and consult a healthcare professional if menstrual pain is severe or affecting daily life.",
    "painful or swollen joints": "You might have arthritis, injury, or inflammation. Rest the affected joints, use ice or heat as needed, and consult a healthcare professional if pain is severe or persistent.",
    "painful or swollen feet": "You might have an injury, fluid retention, or circulatory issues. Rest your feet, elevate them, and consult a healthcare professional if symptoms persist or worsen.",
    "persistent cough or sore throat": "You might have a respiratory infection, allergies, or other issues. Stay hydrated, use cough remedies, and consult a healthcare professional if symptoms are severe or persist for more than a week.",
    "painful or swollen throat": "You might have a throat infection, tonsillitis, or other issues. Gargle with warm salt water, stay hydrated, and consult a healthcare professional if symptoms are severe or persistent.",
    "persistent headaches or migraines": "You might have tension headaches, migraines, or other issues. Keep a headache diary, manage stress, and consult a healthcare professional if headaches are frequent or severe.",
    "painful urination": "You might have a urinary tract infection or other urinary issues. Increase fluid intake, use over-the-counter pain relief, and consult a healthcare professional if symptoms are severe or persistent.",
    "painful or swollen feet": "You might have an injury, fluid retention, or circulatory issues. Rest your feet, elevate them, and consult a healthcare professional if symptoms persist or worsen.",
    "painful or swollen abdomen": "You might have digestive issues, inflammation, or other conditions. Rest, avoid heavy foods, and consult a healthcare professional if symptoms are severe or persistent.",
    "chronic dry skin": "You might have a skin condition or be experiencing environmental dryness. Use moisturizers, stay hydrated, and consult a dermatologist if dryness is severe or persistent.",
    "painful or swollen throat": "You might have a throat infection, tonsillitis, or other issues. Gargle with warm salt water, stay hydrated, and consult a healthcare professional if symptoms are severe or persistent.",
    "persistent headaches or migraines": "You might have tension headaches, migraines, or other issues. Keep a headache diary, manage stress, and consult a healthcare professional if headaches are frequent or severe.",
    "painful urination": "You might have a urinary tract infection or other urinary issues. Increase fluid intake, use over-the-counter pain relief, and consult a healthcare professional if symptoms are severe or persistent.",
    "painful or swollen feet": "You might have an injury, fluid retention, or circulatory issues. Rest your feet, elevate them, and consult a healthcare professional if symptoms persist or worsen.",
    "painful or swollen abdomen": "You might have digestive issues, inflammation, or other conditions. Rest, avoid heavy foods, and consult a healthcare professional if symptoms are severe or persistent.",
    "chronic dry skin": "You might have a skin condition or be experiencing environmental dryness. Use moisturizers, stay hydrated, and consult a dermatologist if dryness is severe or persistent.",
    "painful or swollen throat": "You might have a throat infection, tonsillitis, or other issues. Gargle with warm salt water, stay hydrated, and consult a healthcare professional if symptoms are severe or persistent.",
    "persistent headache or migraine": "You might be experiencing a migraine, tension headache, or other condition. Keep a headache diary, manage stress, and consult a healthcare professional if headaches are severe or frequent.",
    "painful urination": "You might have a urinary tract infection or other urinary issues. Increase fluid intake, use over-the-counter pain relief, and consult a healthcare professional if symptoms are severe or persistent.",
    "painful or swollen feet": "You might have an injury, fluid retention, or circulatory issues. Rest your feet, elevate them, and consult a healthcare professional if symptoms persist or worsen.",
    "painful or swollen abdomen": "You might have digestive issues, inflammation, or other conditions. Rest, avoid heavy foods, and consult a healthcare professional if symptoms are severe or persistent.",
    "persistent dry skin": "You might have a skin condition or be experiencing environmental dryness. Use moisturizers, stay hydrated, and consult a dermatologist if dryness is severe or persistent.",
    "painful or swollen throat": "You might have a throat infection, tonsillitis, or other issues. Gargle with warm salt water, stay hydrated, and consult a healthcare professional if symptoms are severe or persistent.",
    "painful urination or frequent urination": "You might have a urinary tract infection or other issues. Increase fluid intake, use pain relief as needed, and consult a healthcare professional if symptoms are severe or persistent.",
    "painful menstruation": "You might have dysmenorrhea or other menstrual issues. Use pain relievers, apply heat to the abdomen, and consult a healthcare professional if menstrual pain is severe or affecting daily life.",
    "painful or swollen joints": "You might have arthritis, injury, or inflammation. Rest the affected joints, use ice or heat as needed, and consult a healthcare professional if pain is severe or persistent.",
    "persistent or severe headache": "You might be experiencing a migraine, tension headache, or other condition. Keep a headache diary, manage stress, and consult a healthcare professional if headaches are severe or frequent.",
    "painful urination": "You might have a urinary tract infection or other urinary issues. Increase fluid intake, use over-the-counter pain relief, and consult a healthcare professional if symptoms are severe or persistent."
}
# Simulated database of doctors
def fetch_doctors_from_search(location):
    query = f"doctors in {location}"
    search_results = search(query, num_results=10)
    doctor_names = []

    for result in search_results:
        try:
            response = requests.get(result)
            soup = BeautifulSoup(response.content, 'lxml')

            # Extract doctor name based on common HTML patterns
            titles = soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6'])
            for title in titles:
                text = title.get_text()
                if 'Dr.' in text:
                    doctor_names.append(text.strip())
        except Exception as e:
            st.write(f"Error scraping {result}: {e}")

    return doctor_names

doctor_names = [
    "Dr. Aparna", "Dr. Balamurugan", "Dr. Clinton", "Dr. Devika", "Dr. Elias",
    "Dr. Farhan", "Dr. Gitanjali", "Dr. Harish", "Dr. Isha", "Dr. Jayant",
    "Dr. Kavita", "Dr. Lalit", "Dr. Manisha", "Dr. Nitin", "Dr. Ojas",
    "Dr. Pranav", "Dr. Qadir", "Dr. Radhika", "Dr. Sameer", "Dr. Tanvi"
]
# Define health tips based on BMI categories
def get_bmi_tips(bmi):
    if bmi < 18.5:
        return (
            "🔍 **Underweight**\n\n"
            "You are underweight. It's crucial to address this by adopting healthy eating habits and consulting a healthcare provider for personalized advice.\n\n"
            "🍎 **Dietary Tips:**\n"
            "- **Increase Caloric Intake**: Incorporate calorie-dense foods such as nuts, seeds, avocados, and whole grains. 🌰🥑\n"
            "- **Balanced Meals**: Ensure your meals have a good mix of proteins, carbs, and fats. Lean meats, legumes, and dairy are great choices. 🥩🍗\n"
            "- **Frequent Meals**: Eat small, frequent meals throughout the day to boost calorie intake. 🍽️\n"
            "- **Protein-Rich Foods**: Focus on high-protein foods to build muscle mass. 💪🍗\n\n"
            "🏋️‍♂️ **Lifestyle Tips:**\n"
            "- **Strength Training**: Engage in exercises that help build muscle mass. 🏋️‍♂️\n"
            "- **Healthy Snacking**: Include nutritious snacks between meals. 🍎🥜\n"
            "- **Monitor Weight**: Regularly track your weight and dietary intake. 📉\n\n"
            "🩺 **Professional Advice:**\n"
            "- **Consult a Healthcare Provider**: Consult with a healthcare professional to rule out underlying health issues and get tailored dietary recommendations. 🩺"
        )
    elif 18.5 <= bmi < 24.9:
        return (
            "🏅 **Normal Weight**\n\n"
            "You have a normal weight. To maintain your health, continue following a balanced diet and regular physical activity.\n\n"
            "🥗 **Dietary Tips:**\n"
            "- **Balanced Diet**: Continue with a variety of fruits, vegetables, whole grains, lean proteins, and healthy fats. 🍇🥦🍗\n"
            "- **Portion Control**: Be mindful of portion sizes to maintain your weight. 📏\n"
            "- **Hydration**: Drink plenty of water throughout the day. 💧\n\n"
            "🏃‍♂️ **Lifestyle Tips:**\n"
            "- **Regular Exercise**: Maintain a routine that includes cardio, strength, and flexibility exercises. 🏃‍♀️🏋️‍♂️\n"
            "- **Healthy Habits**: Practice mindful eating and avoid excessive sugar and processed foods. 🍴🚫\n"
            "- **Monitor Health**: Regularly check your weight and overall health. 📊\n\n"
            "🩺 **Professional Advice:**\n"
            "- **Routine Check-ups**: Continue with regular health check-ups to ensure you stay in good health. 🩺"
        )
    elif 25 <= bmi < 29.9:
        return (
            "⚠️ **Overweight**\n\n"
            "You are overweight. Adopting a healthier diet and increasing physical activity can help. Consult a healthcare provider if necessary.\n\n"
            "🥦 **Dietary Tips:**\n"
            "- **Reduce Caloric Intake**: Focus on cutting down calorie consumption by avoiding high-calorie foods. 🍔❌\n"
            "- **Increase Fiber**: Eat high-fiber foods like vegetables, fruits, and whole grains to help you feel full. 🥕🍎\n"
            "- **Healthy Choices**: Use healthy cooking methods like grilling and baking instead of frying. 🍴🍳\n\n"
            "🏃 **Lifestyle Tips:**\n"
            "- **Increase Physical Activity**: Aim for at least 150 minutes of moderate aerobic activity per week, plus strength training. 🚴‍♂️🏋️‍♀️\n"
            "- **Behavioral Changes**: Practice mindful eating and avoid emotional eating. 🧠🍽️\n"
            "- **Regular Monitoring**: Track your weight and dietary changes to stay on course. 📉\n\n"
            "🩺 **Professional Advice:**\n"
            "- **Consult a Healthcare Provider**: Seek advice from a dietitian or healthcare provider to create a personalized weight management plan. 🩺"
        )
    else:
        return (
            "🚨 **Obese**\n\n"
            "You are obese. It's important to consult with a healthcare provider for a tailored weight management plan and consider healthier lifestyle habits.\n\n"
            "🥗 **Dietary Tips:**\n"
            "- **Structured Diet Plan**: Follow a controlled diet plan to reduce calorie intake while maintaining balanced nutrition. 🍽️📋\n"
            "- **Avoid High-Sugar Foods**: Minimize sugary foods and drinks that contribute to weight gain. 🍭❌\n"
            "- **Nutrient-Dense Foods**: Focus on consuming nutrient-rich foods like vegetables, fruits, lean proteins, and whole grains. 🥦🍎🍗\n\n"
            "🏋️‍♂️ **Lifestyle Tips:**\n"
            "- **Increase Physical Activity**: Engage in at least 150 minutes of moderate-intensity aerobic exercise per week, including strength training. 🏃‍♀️🏋️\n"
            "- **Behavioral Therapy**: Consider therapy to address any eating habits or emotional factors contributing to obesity. 🧠💬\n"
            "- **Regular Monitoring**: Keep track of weight loss progress and adjust your plan as needed. 📈\n\n"
            "🩺 **Professional Advice:**\n"
            "- **Consult a Healthcare Provider**: Consult with a healthcare provider to develop a comprehensive weight management plan, including potential medical interventions. 🩺"
        )


# Function to get chat completion from Groq API
def get_chat_completion(messages):
    api_key = os.environ.get("GROQ_API_KEY")
    if not api_key:
        return "API key is not set. Please check the environment variable."
    client = ChatGroq(api_key=api_key, model='llama3-70b-8192')
    try:
        response = client.invoke(messages)
        return response.content
    except Exception as e:
        return f"An error occurred: {e}"

# Function to diagnose based on training data or API response
def diagnose(symptoms):
    symptom_lower = symptoms.lower()
    if symptom_lower in training_data:
        return training_data[symptom_lower]
    else:
        # Use the API if the symptom is not in the training data
        messages = [SystemMessage("You are a virtual doctor that you predict the diesease name and give the medical assesments and remedies ."), HumanMessage(f"I have {symptoms}. What should I do?")]
        response = get_chat_completion(messages)
        return response

# Function to check reminders
def check_reminders():
    current_time = datetime.now().strftime("%H:%M")
    reminders = st.session_state.get('reminders', [])
    for reminder in reminders:
        if reminder['time'] == current_time:
            st.session_state['alert'] = True
            st.session_state['alert_message'] = f"Time to take your medication: {reminder['medication']}"
            play_alarm()
            break

# Function to play alarm sound
def play_alarm():
    st.write("""
        <audio autoplay>
            <source src="https://www.soundjay.com/button/sounds/button-4.mp3" type="audio/mpeg">
        </audio>
    """, unsafe_allow_html=True)

# Sidebar content with radio buttons

st.markdown(
    """
    <style>
    /* Increase the size of the radio buttons */
    input[type="radio"] {
        transform: scale(1.5);  /* Adjust the scale as needed */
        margin-right: 10px;  /* Adjust the margin as needed */
    }
    </style>
    """,
    unsafe_allow_html=True
)
def sidebar():
    st.sidebar.title("Information")
    option = st.sidebar.selectbox(
    "***Go Deeper through the User Interface***",
    ["Overview", "Features", "FAQ"]
)


    if option == "Overview":
        st.sidebar.markdown(
        """
        🌟 Welcome to Your Health Revolution! 🌟

With our Virtual Doctor application, you get:

📜 Instant Symptom Diagnosis: Describe your symptoms and get immediate, actionable advice.

📅 Doctor Appointments: Schedule meetings with specialists and find nearby doctors effortlessly.

📝 Daily Health Tips: Receive daily insights to boost your well-being.

⏰ Medication Reminders: Never miss a dose with personalized reminders.

📊 BMI Calculator: Assess your body mass index and get tailored health advice.

📂 Export Chat History: Save and review your health conversations for future reference.

Transform your healthcare experience with ease and convenience. Embrace a healthier you today!
        """,
        unsafe_allow_html=True
    )

    elif option == "Features":
        st.sidebar.markdown(
        """
        ## Features

        - **🔍 Symptom Diagnosis and Remedies**: Discover the power of an AI-driven virtual doctor. Simply describe your symptoms to receive accurate diagnoses and personalized remedies tailored just for you.

        - **📜 Chat History Export**: Keep track of your health journey by saving your chat history to a text file. Perfect for reviewing past interactions and sharing with healthcare providers.

        - **📅 Schedule a Meeting**: Easily book appointments with top specialized doctors. Plus, find the nearest medical professionals to ensure you get the care you need promptly.

        - **🔐 Authentication**: Experience a secure and private environment with our robust login system, protecting your personal health information.

        - **🌟 Daily Health Tips**: Stay informed with fresh, daily health tips designed to help you maintain a balanced lifestyle and make healthier choices.

        - **⏰ Medication Reminders**: Never miss a dose again! Set up medication reminders to receive timely alerts and keep your health on track.

        - **⚖️ BMI Calculator**: Instantly calculate your Body Mass Index and receive insightful, personalized health advice based on your results.

        - **📞 Emergency Contact Information**: Quickly access essential emergency contact details for peace of mind and readiness in critical situations.

        Get ready to transform your health management experience with these innovative features!
        """,
        unsafe_allow_html=True
    )


    elif option == "FAQ":
        st.sidebar.markdown(
        """
        ## FAQ

        **Q: How do I use this application?**  
        A: Log in with your credentials. Enter your symptoms to get a diagnosis or schedule a meeting with a doctor.

        **Q: What should I do if my symptom is not listed?**  
        A: The virtual doctor will provide responses based on the available data even if your symptom is not listed.

        **Q: How can I export my chat history?**  
        A: Click the 'Export Chat History' button to save your chat history to a text file.

        **Q: How do I schedule a meeting with a doctor?**  
        A: Go to the 'Book A DOC' tab and choose from available options.

        **Q: What types of doctors can I schedule a meeting with?**  
        A: You can schedule appointments with various specialists such as cardiologists, dermatologists, pediatricians, and more.

        **Q: How do I set a medication reminder?**  
        A: Navigate to the 'Medication Reminder' tab, enter the medication details and the reminder time, and save it.

        **Q: How is my BMI calculated?**  
        A: Your BMI is calculated based on your height and weight using the formula: weight (kg) / (height (m)²).

        **Q: How do I update my profile or change my password?**  
        A: Profile updates and password changes can be done through the 'Profile' section of the application settings.

        **Q: Can I access this application on mobile devices?**  
        A: Yes, this application is accessible on most mobile devices and tablets via a web browser.

        **Q: Who should I contact for technical support or issues?**  
        A: For any technical support or issues, please contact our support team via the 'Contact Us' section or email support@HealthMate.com.
        """,
        unsafe_allow_html=True
    )


    # Add a small logo or image at the bottom of the sidebar
    st.sidebar.markdown(
        """
        <style>
        .css-1v3fvcr {
            display: flex;
            flex-direction: column;
            height: 100%;
            justify-content: space-between;
        }
        .css-1v3fvcr img {
            margin-top: auto;
            max-width: 20px; /* Adjust size here */
            max-height: 20px; /* Adjust size here */
        }
        </style>
        """,
        unsafe_allow_html=True
    )

    #st.sidebar.image(r"C:\Users\praje\OneDrive\AppData\Desktop\project\health\work\assets\logo-no-background.png", use_column_width=False)

def create_map(location):
    geolocator = Nominatim(user_agent="healthmate")
    geocode_result = geolocator.geocode(location)

    if geocode_result:
        latitude = geocode_result.latitude
        longitude = geocode_result.longitude
        m = folium.Map(location=[latitude, longitude], zoom_start=13)
        folium.Marker([latitude, longitude], popup="Your Location", tooltip="Your Location").add_to(m)
        return m
    else:
        st.write("🚫 Location not found. Please enter a valid location.")
        return None





# Streamlit interface
def main():
    sidebar()
    st.title("**🩺 HealthMate**")
    #increase the font size

    #st.write("Welcome to Your Virtual Doctor! Describe your symptoms and get a diagnosis and remedy.")  # Description
    

    #sidebar()

    # Check authentication status
    if 'messages' not in st.session_state:
        st.session_state.messages = []
    if 'diagnosis' not in st.session_state:
        st.session_state.diagnosis = ""
    if 'all_symptoms' not in st.session_state:
        st.session_state.all_symptoms = ""
    if 'authenticated' not in st.session_state:
        st.session_state.authenticated = False
    if 'map' not in st.session_state:
        st.session_state.map = None
    if 'doctors' not in st.session_state:
        st.session_state.doctors = []
    if 'location' not in st.session_state:
        st.session_state.location = ""

    if not st.session_state.authenticated:
        st.sidebar.header("Login")
        username = st.text_input("Username:")
        password = st.text_input("Password:", type="password")
        
        if st.button("Login"):
            if username == "Prajen SK" and password == "Prajen@2004":
                st.session_state.authenticated = True
                st.session_state.username = username
                st.success("Logged in successfully!")
            else:
                st.session_state.authenticated = False
                st.error("Invalid credentials!")
    else:
        st.session_state['alert'] = False

        # Horizontal tab layout
        tabs = st.tabs(["Virtual DOC", "Book a DOC", " Medication Reminder", "BMI Calculator", "Contact Emergency", "Daily health tips"])
        
        with tabs[0]:
            st.header(" I am Your Virtual Doctor")
            st.write("✏️ Describe your symptoms and get a diagnosis and remedy.")

            if 'messages' not in st.session_state:
                st.session_state.messages = []
            if 'diagnosis' not in st.session_state:
                st.session_state.diagnosis = ""

            symptoms = st.text_input("🔍 Enter your symptoms:")

            if st.button("🏥 Get Diagnosis"):
                if symptoms:
                    st.session_state.all_symptoms += " " + symptoms
                    combined_symptoms = st.session_state.all_symptoms.strip()
                    diagnosis = diagnose(combined_symptoms)
                    st.session_state.messages.append({"role": "user", "content": symptoms})
                    st.session_state.messages.append({"role": "assistant", "content": diagnosis})
                    st.session_state.diagnosis = diagnosis
                else:
                    st.write("Please enter your symptoms.")
            if st.session_state.diagnosis:
                st.write(f"**Diagnosis:** {st.session_state.diagnosis}")


            if st.button("📖 Export Chat History"):
                if st.session_state.messages:
                    try:
                # Load the template
                        doc = Document("template.docx")

                # Add chat history to the document
                        doc.add_paragraph("Chat History", style='Heading1')
                        for message in st.session_state.messages:
                            role, content = message['role'], message['content']
                            doc.add_paragraph(f"{role.capitalize()}: {content}")

                # Save the document to an in-memory file
                        buffer = io.BytesIO()
                        doc.save(buffer)
                        buffer.seek(0)

                # Provide download button
                        st.download_button(
                            label="Download Chat History",
                            data=buffer,
                            file_name="chat_history.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )

                        st.success("Chat history is ready for download.")
                    except Exception as e:
                        st.error(f"Error writing to file: {e}")
            else:
                st.write("No chat history to export.")

        with tabs[1]:
            st.header("Meet a Doctor for Further Consultation")
            st.write("📅 **Schedule a meeting with a specialized doctor**")

            location = st.text_input("📍 Enter your location to find nearby doctors:")

            if st.button("🔍 Check Nearby Doctors"):
                if location:
                    st.session_state.map = create_map(location)
                    st.session_state.location = location

                    st.write(f"🔎 Checking for doctors in {location}...")
                    st.session_state.doctors = fetch_doctors_from_search(location)
                else:
                    st.write("🚫 Please enter your location.")

            if st.session_state.doctors:
                st.subheader("👨‍⚕️ Available Doctors:")
                for doctor in st.session_state.doctors:
                    st.write(f"- {doctor} (Location: {st.session_state.location})")

            if st.session_state.map:
                st_folium(st.session_state.map, width=700, height=500, use_container_width=True)

            st.subheader("📅 Request an Appointment")
            name = st.text_input("📝 Your Name")
            contact = st.text_input("📞 Contact Number")
            if st.button("📩 Submit Request"):
                if name and contact:
                    st.write("✅ Your appointment request has been submitted to the doctors. They will get back to you soon.")
                else:
                    st.write("🚫 Please fill in all required fields.")

        with tabs[2]:
            st.header("⏰ Medication Reminder")
            st.write("Set Medication Reminders")
            medication = st.text_input("Medication Name")
            reminder_time = st.time_input("Reminder Time")

            if st.button("Set Reminder ⏳"):
                if medication and reminder_time:
                    reminders = st.session_state.get('reminders', [])
                    reminders.append({"medication": medication, "time": reminder_time.strftime("%H:%M")})
                    st.session_state['reminders'] = reminders
                    st.write(f"Reminder set for {medication} at {reminder_time.strftime('%H:%M')}")
                else:
                    st.write("Please enter medication name and time.")
            
        with tabs[3]:
            st.header("🧮 BMI Calculator")
            st.write("## Calculate your Body Mass Index (BMI)")
            weight = st.number_input("Enter your weight (kg):", min_value=0.0)
            height = st.number_input("Enter your height (cm):", min_value=0.0)

            if st.button("🧮 Calculate BMI"):
                if weight > 0 and height > 0:
                    height_m = height / 100
                    bmi = weight / (height_m ** 2)
                    st.write(f"📊 Your BMI is: {bmi:.2f}")
                    
                    # Provide health condition and tips
                    health_tips = get_bmi_tips(bmi)
                    st.write(f"**Health Condition:** {health_tips}")
                else:
                    st.write("Please enter valid weight and height.")

        
        with tabs[4]:
            st.header("📞 Contact Emergency")
            st.write("## Emergency Contact Information")

            st.write("""
 In Tamil Nadu, it is crucial to have access to reliable emergency services. Here are some essential contacts you should keep handy:

    - **🚑 Ambulance Services**:  
      For immediate medical transport, call **108**. This service provides ambulances equipped with life-saving equipment and personnel trained to handle critical medical situations. Available 24/7 across Tamil Nadu, this number ensures you get prompt medical attention when you need it the most.

    - **🚓 Police Services**:  
      In case of any law enforcement or safety emergencies, dial **100**. This number connects you to local police stations throughout Tamil Nadu, ensuring rapid response to emergencies such as theft, assault, or other criminal activities.

    - **🚒 Fire Services**:  
      For emergencies involving fire or hazardous situations, call **101**. Tamil Nadu Fire and Rescue Services provide swift assistance for fire incidents, including rescue operations and firefighting. Their teams are trained to handle various fire-related emergencies.

    - **🏥 Emergency Hospitals**:  
      In case you need immediate medical care, contact these major hospitals:
      - **Chennai**: Government General Hospital  
        📞 Phone: +91 44 2522 9555  
        📍 Address: Park Town, Chennai, Tamil Nadu, 600003  
        Provides comprehensive emergency services including trauma care, surgery, and more.
        
      - **Coimbatore**: Coimbatore Medical College Hospital  
        📞 Phone: +91 422 245 0422  
        📍 Address: Coimbatore, Tamil Nadu, 641018  
        Offers emergency medical treatment, including specialized care and diagnostics.
        
      - **Madurai**: Government Rajaji Hospital  
        📞 Phone: +91 452 253 4023  
        📍 Address: Madurai, Tamil Nadu, 625002  
        Equipped to handle various emergencies and critical care situations.

    - **☎️ Toll-Free Helplines**:
      - **Women's Helpline**: **181**  
        Dedicated to providing assistance and support to women in distress, this helpline offers counseling and emergency support for issues like domestic violence, harassment, and more.
        
      - **Child Helpline**: **1098**  
        Focused on child protection and welfare, this helpline provides support for children in need, including cases of abuse, neglect, and other child-related emergencies.

    In an emergency, ensure you provide clear information about your location and the nature of the emergency to get the quickest and most appropriate help. Remember, these services are here to assist you around the clock, so don't hesitate to reach out when needed. Your safety is our utmost priority.
    """)


                    
        with tabs[5]:
            #st.header("🩺 Health Tips")
            st.header(" Daily Health Tips")

            st.write("**1. 💧 Stay Hydrated:** Drink at least 8 glasses of water a day. Hydration is crucial for maintaining bodily functions, improving skin health, and supporting digestion.")

            st.write("**2. 🏃 Exercise Regularly:** Aim for at least 150 minutes of moderate aerobic activity or 75 minutes of vigorous activity per week. Incorporate strength training exercises twice a week to build and maintain muscle mass.")

            st.write("**3. 💤 Get Enough Sleep:** Strive for 7-9 hours of quality sleep each night. Good sleep is essential for recovery, mental health, and overall well-being.")

            st.write("**4. 🥗 Eat a Balanced Diet:** Include a variety of fruits, vegetables, lean proteins, and whole grains in your meals. Limit intake of processed foods, sugars, and saturated fats.")

            st.write("**5. ⚖️ Maintain a Healthy Weight:** Balance your calorie intake with physical activity to keep your weight in check. Consult a healthcare provider for personalized advice.")

            st.write("**6. 🧼 Practice Good Hygiene:** Wash your hands regularly, brush and floss your teeth twice a day, and keep your living environment clean to prevent illnesses.")

            st.write("**7. 🧘 Manage Stress:** Engage in relaxation techniques such as meditation, deep breathing exercises, or hobbies to manage stress levels effectively.")

            st.write("**8. 🚭 Avoid Smoking and Limit Alcohol:** Smoking and excessive alcohol consumption can lead to various health problems. Seek support if you need help quitting.")

            st.write("**9. 🩺 Regular Health Check-ups:** Schedule routine check-ups with your healthcare provider to monitor your health and catch any potential issues early.")

            st.write("**10. 🤝 Stay Connected:** Maintain social connections and seek support from friends and family to boost mental health and overall happiness.")

            st.write("Incorporate these tips into your daily routine to promote a healthier lifestyle and achieve optimal fitness. Remember, small consistent changes can make a big difference!")

    # You can also add more tips or fetch them from a database or API as needed.



        st_autorefresh(interval=60 * 1000, key="refresh")  # Refresh every minute
        check_reminders()

if __name__ == "__main__":
    main()











#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